# ── Config ────────────────────────────────────────────────────────────────────
MODEL = "claude-haiku-4-5-20251001"
MAX_TOKENS = 4096
OUTPUTS_DIR = "outputs"
# ─────────────────────────────────────────────────────────────────────────────

import json
import os
import re
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from dotenv import load_dotenv

load_dotenv()

_client: anthropic.Anthropic | None = None


def _get_client() -> anthropic.Anthropic:
    global _client
    if _client is None:
        _client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    return _client


_SYSTEM_PROMPT = """\
You are an expert CV tailor and career consultant.

Given a candidate's CV and a job description you will:
1. Extract the 10 most important keywords/phrases from the job description.
2. Rewrite the CV bullet points to naturally incorporate those keywords while preserving factual accuracy — never invent experience the candidate does not have.
3. Write a cover letter following the rules below.
4. Score how well the tailored CV matches the job description on a scale of 0–100.

COVER LETTER RULES — follow every one of these:
- Exactly 3 paragraphs, each 3–4 sentences. Paragraphs separated by a blank line.
- Write in first person throughout.
- Vary sentence length. Mix short punchy sentences with longer ones. Never write three sentences of similar length in a row.
- Sound like a real person wrote this, not a template. Specific beats generic.
- Forbidden words — never use any of these: leverage, spearhead, delve, foster, streamline, robust, dynamic, passionate, synergy, transformative, impactful
- Forbidden phrases — never write: "proven track record", "results-driven", "results-oriented", "team player", "go-getter", "self-starter", "detail-oriented", "think outside the box", "hit the ground running", "wear many hats"
- Confident but not boastful. State facts about what you did; let the work speak.
- Reference the company or role specifically — no generic "your esteemed organisation".
- Paragraph 1: why this specific role at this specific company, grounded in something real from the job description.
- Paragraph 2: one or two concrete things from the CV that are directly relevant — be specific, use numbers if available.
- Paragraph 3: short close. Express genuine interest, state you'd welcome a conversation. No hollow sign-offs.

Respond ONLY with valid JSON — no markdown fences, no preamble — in this exact shape:
{
  "keywords": ["kw1", "kw2", "kw3", "kw4", "kw5", "kw6", "kw7", "kw8", "kw9", "kw10"],
  "tailored_cv": "<full tailored CV as plain text>",
  "cover_letter": "<3-paragraph cover letter, paragraphs separated by blank lines>",
  "match_score": <integer 0-100>
}
"""


def _parse_json(text: str) -> dict:
    text = text.strip()
    # Strip optional markdown code fence if the model includes one anyway
    fence = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if fence:
        text = fence.group(1).strip()
    return json.loads(text)


def _save_to_docx(tailored_cv: str, cover_letter: str) -> Path:
    doc = Document()

    doc.add_heading("Tailored CV", level=1)
    for line in tailored_cv.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_page_break()

    doc.add_heading("Cover Letter", level=1)
    for paragraph in cover_letter.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    out_dir = Path(OUTPUTS_DIR)
    out_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = out_dir / f"tailored_cv_{timestamp}.docx"
    doc.save(path)
    return path


def tailor_cv(cv_text: str, job_description: str) -> dict:
    """Tailor a CV to a job description using Claude.

    Returns a dict with keys:
        keywords    – list[str]   top 10 JD keywords
        tailored_cv – str         rewritten CV plain text
        cover_letter– str         3-paragraph cover letter
        match_score – int         0–100 quality score
        output_path – str         path to the saved .docx file
    """
    client = _get_client()

    response = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=[
            {
                "type": "text",
                "text": _SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[
            {
                "role": "user",
                "content": (
                    f"CV:\n{cv_text}\n\n"
                    f"Job Description:\n{job_description}"
                ),
            }
        ],
    )

    result = _parse_json(response.content[0].text)

    required = {"keywords", "tailored_cv", "cover_letter", "match_score"}
    missing = required - result.keys()
    if missing:
        raise ValueError(f"Claude response missing fields: {missing}")

    output_path = _save_to_docx(result["tailored_cv"], result["cover_letter"])
    result["output_path"] = str(output_path)
    return result
